import streamlit as st
import pdfplumber
import time
import re
import os
import pytesseract
import openai
from docx import Document
from io import BytesIO
from deep_translator import GoogleTranslator
from gtts import gTTS
from PIL import Image
from pdf2image import convert_from_bytes

# ---------------- CONFIGURAÇÃO ----------------
st.set_page_config(page_title="DocuTools Pro", page_icon="🚀", layout="wide")

try:
    API_KEY = st.secrets.get("OPENAI_API_KEY", None)
except Exception:
    API_KEY = None
if not API_KEY:
    API_KEY = os.environ.get("OPENAI_API_KEY", None)

MAX_MB = 50
MAX_OCR_PAGES = 25

IDIOMAS = {
    "Auto-Detectar": "auto",
    "Português": "pt",
    "Inglês": "en",
    "Espanhol": "es",
    "Francês": "fr",
    "Alemão": "de",
    "Italiano": "it",
}

VOZES = {"Português": "pt", "Inglês": "en", "Espanhol": "es", "Francês": "fr"}

FORMATOS_IMG = {"PNG": "PNG", "JPEG": "JPEG", "WEBP": "WEBP", "BMP": "BMP"}

for k in ("resultado", "nome_arquivo", "avisos", "img_processada",
          "img_nome", "audio_mp3", "pdf_gerado"):
    st.session_state.setdefault(k, None)

# ---------------- FUNÇÕES ----------------

def call_with_retry(func, *args, retries=3, delay=1, **kwargs):
    for i in range(retries):
        try:
            return func(*args, **kwargs)
        except Exception:
            if i == retries - 1:
                return None
            time.sleep(delay * (2 ** i))
    return None


def tem_conteudo_traduzivel(linha):
    return bool(re.search(r"[A-Za-zÀ-ÿ]", linha))


@st.cache_data(show_spinner=False)
def extrair_texto(file_bytes, file_type):
    try:
        if "pdf" in file_type:
            with pdfplumber.open(BytesIO(file_bytes)) as pdf:
                texto = "\n".join([p.extract_text() or "" for p in pdf.pages])
            if len(texto.strip()) < 20:
                imagens = convert_from_bytes(
                    file_bytes, dpi=200, first_page=1, last_page=MAX_OCR_PAGES
                )
                texto = "\n".join(
                    [pytesseract.image_to_string(img, lang="por+eng") for img in imagens]
                )
            return texto, None
        elif "officedocument" in file_type:
            doc = Document(BytesIO(file_bytes))
            return "\n".join([p.text for p in doc.paragraphs]), None
        elif "image" in file_type:
            img = Image.open(BytesIO(file_bytes))
            return pytesseract.image_to_string(img, lang="por+eng"), None
        else:
            try:
                return file_bytes.decode("utf-8"), None
            except UnicodeDecodeError:
                return file_bytes.decode("latin-1"), None
    except Exception as e:
        return None, str(e)


def traduzir_openai(texto, de, para):
    client = openai.OpenAI(api_key=API_KEY)
    res = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {
                "role": "system",
                "content": f"Traduza de {de} para {para}. Preserve termos técnicos "
                f"e a estrutura. Retorne APENAS o texto traduzido.",
            },
            {"role": "user", "content": texto},
        ],
        temperature=0.2,
    )
    return res.choices[0].message.content


def traduzir_bloco(texto, de, para):
    """NUNCA retorna None. No pior caso, devolve o texto original."""
    if not texto or not texto.strip() or not tem_conteudo_traduzivel(texto):
        return texto or ""

    if API_KEY:
        resposta = call_with_retry(traduzir_openai, texto, de, para)
        if isinstance(resposta, str) and resposta.strip():
            return resposta

    origem = IDIOMAS.get(de, "auto")
    destino = IDIOMAS.get(para, "pt")
    resultado = call_with_retry(
        GoogleTranslator(source=origem, target=destino).translate, texto
    )
    return resultado if isinstance(resultado, str) and resultado.strip() else texto


def traduzir_documento(texto, de, para):
    linhas = texto.split("\n")
    if not linhas:
        return "", 0

    final, falhas = [], 0
    barra = st.progress(0, text="Traduzindo parágrafos...")
    for i, linha in enumerate(linhas):
        if linha.strip():
            t = traduzir_bloco(linha, de, para)
            if t == linha and tem_conteudo_traduzivel(linha) and de != para:
                falhas += 1
            final.append(t if isinstance(t, str) else linha)
        else:
            final.append("")
        barra.progress((i + 1) / len(linhas))
    barra.empty()
    return "\n".join(final), falhas


def gerar_docx(texto):
    doc = Document()
    for linha in texto.split("\n"):
        doc.add_paragraph(linha)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------- INTERFACE ----------------

st.title("🚀 DocuTools Pro")
st.caption("Processamento inteligente de documentos — Tradução, OCR, Imagem e Áudio")

tabs = st.tabs(["🌐 Tradução & OCR", "🖼️ Ferramentas de Imagem", "🔊 Áudio", "⚙️ Sistema"])

# ---- ABA 1: TRADUÇÃO ----
with tabs[0]:
    st.subheader("Tradução com preservação de estrutura")
    arq = st.file_uploader(
        "Envie PDF, DOCX, TXT ou Imagem",
        type=["pdf", "docx", "txt", "png", "jpg", "jpeg"],
        key="up_doc",
    )

    col1, col2 = st.columns(2)
    with col1:
        de_lang = st.selectbox("Idioma original", list(IDIOMAS.keys()), index=0)
    with col2:
        destinos = [k for k in IDIOMAS if k != "Auto-Detectar"]
        para_lang = st.selectbox("Traduzir para", destinos, index=0)

    if arq:
        if arq.size > MAX_MB * 1024 * 1024:
            st.error(f"Arquivo excede {MAX_MB} MB.")
            st.stop()

        if st.button("🚀 Traduzir Documento", type="primary", key="btn_trad"):
            with st.status("Processando...", expanded=True) as status:
                st.write("📖 Extraindo texto (OCR automático se necessário)...")
                texto, erro = extrair_texto(arq.getvalue(), arq.type)

                if erro:
                    status.update(label="Falha na extração", state="error")
                    st.error(f"Erro ao ler o arquivo: {erro}")
                    st.stop()
                if not texto or not texto.strip():
                    status.update(label="Documento vazio", state="error")
                    st.warning("Nenhum texto encontrado no documento.")
                    st.stop()

                st.write("🌐 Traduzindo parágrafo por parágrafo...")
                resultado, falhas = traduzir_documento(texto, de_lang, para_lang)

                st.session_state.resultado = resultado
                st.session_state.nome_arquivo = arq.name
                st.session_state.avisos = falhas
                status.update(label="Concluído!", state="complete")

    if st.session_state.resultado:
        if st.session_state.avisos:
            st.warning(
                f"{st.session_state.avisos} trecho(s) mantidos no idioma original "
                f"(instabilidade momentânea do tradutor)."
            )
        st.text_area("Resultado:", st.session_state.resultado, height=350)

        c1, c2 = st.columns(2)
        with c1:
            st.download_button(
                "📥 Baixar TXT",
                st.session_state.resultado,
                file_name=f"DocuTools_{st.session_state.nome_arquivo}.txt",
                key="dl_txt",
            )
        with c2:
            st.download_button(
                "📥 Baixar DOCX",
                gerar_docx(st.session_state.resultado),
                file_name=f"DocuTools_{st.session_state.nome_arquivo}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="dl_docx",
            )

# ---- ABA 2: FERRAMENTAS DE IMAGEM (leves, sem IA) ----
with tabs[1]:
    st.subheader("Redimensionar, converter e gerar PDF")

    ferramenta = st.radio(
        "Escolha a ferramenta:",
        ["📐 Redimensionar", "🔄 Converter formato", "📑 Imagens → PDF"],
        horizontal=True,
        key="img_tool",
    )

    if ferramenta == "📐 Redimensionar":
        img_arq = st.file_uploader("Envie uma imagem", type=["jpg", "png", "jpeg", "webp", "bmp"], key="up_resize")
        if img_arq:
            img = Image.open(img_arq)
            st.caption(f"Tamanho original: {img.width} × {img.height} px")
            c1, c2 = st.columns(2)
            nova_larg = c1.number_input("Largura (px)", 1, 10000, img.width)
            nova_alt = c2.number_input("Altura (px)", 1, 10000, img.height)
            manter = st.checkbox("Manter proporção", value=True)

            if st.button("Redimensionar", key="btn_resize"):
                if manter:
                    ratio = nova_larg / img.width
                    nova_alt = int(img.height * ratio)
                redim = img.resize((int(nova_larg), int(nova_alt)), Image.LANCZOS)
                buf = BytesIO()
                fmt = "PNG" if img_arq.name.lower().endswith("png") else "JPEG"
                if fmt == "JPEG" and redim.mode in ("RGBA", "P"):
                    redim = redim.convert("RGB")
                redim.save(buf, format=fmt)
                st.session_state.img_processada = buf.getvalue()
                st.session_state.img_nome = f"redimensionada.{fmt.lower()}"
                st.image(st.session_state.img_processada, caption=f"{int(nova_larg)} × {int(nova_alt)} px")

    elif ferramenta == "🔄 Converter formato":
        img_arq = st.file_uploader("Envie uma imagem", type=["jpg", "png", "jpeg", "webp", "bmp"], key="up_conv")
        destino_fmt = st.selectbox("Converter para:", list(FORMATOS_IMG.keys()))
        if img_arq and st.button("Converter", key="btn_conv"):
            img = Image.open(img_arq)
            if destino_fmt == "JPEG" and img.mode in ("RGBA", "P"):
                img = img.convert("RGB")
            buf = BytesIO()
            img.save(buf, format=FORMATOS_IMG[destino_fmt])
            st.session_state.img_processada = buf.getvalue()
            st.session_state.img_nome = f"convertida.{destino_fmt.lower()}"
            st.success(f"Convertida para {destino_fmt}!")

    elif ferramenta == "📑 Imagens → PDF":
        imgs = st.file_uploader(
            "Envie uma ou mais imagens (na ordem desejada)",
            type=["jpg", "png", "jpeg"],
            accept_multiple_files=True,
            key="up_pdf",
        )
        if imgs and st.button("Gerar PDF", key="btn_imgpdf"):
            paginas = []
            for f in imgs:
                im = Image.open(f)
                if im.mode in ("RGBA", "P"):
                    im = im.convert("RGB")
                paginas.append(im)
            buf = BytesIO()
            paginas[0].save(buf, format="PDF", save_all=True, append_images=paginas[1:])
            st.session_state.pdf_gerado = buf.getvalue()
            st.success(f"PDF gerado com {len(paginas)} página(s)!")

    if st.session_state.img_processada:
        st.download_button(
            "📥 Baixar imagem",
            st.session_state.img_processada,
            file_name=st.session_state.img_nome or "imagem.png",
            key="dl_img",
        )
    if st.session_state.pdf_gerado:
        st.download_button(
            "📥 Baixar PDF",
            st.session_state.pdf_gerado,
            file_name="DocuTools_imagens.pdf",
            mime="application/pdf",
            key="dl_imgpdf",
        )

# ---- ABA 3: ÁUDIO ----
with tabs[2]:
    st.subheader("Texto para voz (MP3)")
    txt_voz = st.text_area(
        "Texto:", "DocuTools Pro transformando seu fluxo de trabalho.", key="tts_txt"
    )
    voz = st.selectbox("Idioma da voz", list(VOZES.keys()), key="tts_lang")

    if st.button("🔊 Gerar Áudio", key="btn_tts"):
        if not txt_voz.strip():
            st.warning("Digite um texto primeiro.")
        else:
            with st.spinner("Sintetizando voz..."):
                def _gera():
                    tts = gTTS(text=txt_voz, lang=VOZES[voz])
                    b = BytesIO()
                    tts.write_to_fp(b)
                    return b.getvalue()

                audio = call_with_retry(_gera)
            if audio:
                st.session_state.audio_mp3 = audio
            else:
                st.error("Serviço de voz indisponível. Tente em instantes.")

    if st.session_state.audio_mp3:
        st.audio(st.session_state.audio_mp3, format="audio/mp3")
        st.download_button(
            "📥 Baixar MP3",
            st.session_state.audio_mp3,
            "audio_docutools.mp3",
            "audio/mp3",
            key="dl_mp3",
        )

# ---- ABA 4: SISTEMA ----
with tabs[3]:
    st.subheader("Status do sistema")
    c1, c2, c3 = st.columns(3)
    c1.metric("Motor de Tradução", "GPT-4o-mini" if API_KEY else "Google (grátis)")
    c2.metric("OCR", "Tesseract (por/eng)")
    c3.metric("Limite de upload", f"{MAX_MB} MB")

    if not API_KEY:
        st.info(
            "💡 Configure OPENAI_API_KEY nas variáveis de ambiente "
            "para ativar tradução premium."
        )

    if st.button("🧹 Limpar cache e resultados", key="btn_reset"):
        st.cache_data.clear()
        for k in ("resultado", "nome_arquivo", "avisos", "img_processada",
                  "img_nome", "audio_mp3", "pdf_gerado"):
            st.session_state[k] = None
        st.rerun()

st.sidebar.markdown("### 🚀 DocuTools Pro")
st.sidebar.caption("v3.1 — Leve e estável")
st.sidebar.markdown("---")
st.sidebar.write("✅ Tradução resiliente com retry")
st.sidebar.write("✅ OCR automático (por/eng)")
st.sidebar.write("✅ Imagens: redimensionar, converter, PDF")
st.sidebar.write("✅ Roda em 512 MB de RAM")
